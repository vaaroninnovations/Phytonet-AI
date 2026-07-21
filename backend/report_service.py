"""AI Scientific Report generator using Emergent LLM key (Claude Sonnet 4.5).

Given a full workflow payload (plant, disease, compounds, ADMET, targets,
intersection, PPI, hubs, GO, KEGG, docking, MD), synthesize a publication-ready
IMRAD-style manuscript in Markdown. The report is then converted to HTML / PDF
/ DOCX on demand.
"""
from __future__ import annotations
import asyncio
import io
import json
import logging
import os
import re
import uuid
from datetime import datetime, timezone
from typing import Any, Dict, Optional

from emergentintegrations.llm.chat import LlmChat, UserMessage

logger = logging.getLogger(__name__)

SYSTEM_PROMPT = """You are a senior computational pharmacology researcher writing a
publication-ready manuscript. Voice: formal, third-person, precise. Cite methods
with real tool names & versions. Use SI units. Return valid Markdown."""

SECTION_ORDER = [
    "Title", "Abstract", "Introduction",
    "Materials and Methods",
    "Results",
    "Discussion", "Conclusion", "Limitations", "Future Perspectives",
    "References", "Supplementary Tables"
]


def _kbytes(x, cap: int = 1500) -> str:
    """JSON-encode x, hard-cap at `cap` chars. Used to keep individual data
    blobs from blowing up the LLM context window (Groq llama-3.3-70b is 32k)."""
    try:
        s = json.dumps(x, default=str)
        if len(s) < cap:
            return s
        return s[:cap] + " …truncated"
    except Exception:
        return str(x)[:cap]


def _slim_docking(docking):
    """Reduce docking results to only the fields the LLM needs — dropping the
    per-pose list and heavy interactions dict which can otherwise dominate the
    prompt (10s of KB per pair)."""
    out = []
    for r in (docking or [])[:10]:
        top_hb = [h.get("residue") for h in (r.get("interactions", {}).get("hydrogen_bonds") or [])[:3]]
        top_hp = [h.get("residue") for h in (r.get("interactions", {}).get("hydrophobic_contacts") or [])[:3]]
        cls = r.get("classification") or {}
        out.append({
            "ligand": r.get("ligand_name"),
            "target": r.get("receptor_uniprot"),
            "pdb":    r.get("receptor_pdb"),
            "affinity_kcal_mol": r.get("best_affinity"),
            "quality": cls.get("class"),
            "score":   cls.get("score"),
            "top_hbond_residues":       top_hb,
            "top_hydrophobic_residues": top_hp,
            "n_poses": len(r.get("poses") or []),
        })
    return out


def _build_prompt(workflow: Dict[str, Any]) -> str:
    plant = workflow.get("plant_name") or "the studied plant"
    disease = workflow.get("disease_name") or "the studied disease"
    compounds = workflow.get("selected_compounds") or []
    intersect = workflow.get("intersecting_genes") or []
    hubs = workflow.get("hub_ranking") or []
    go = workflow.get("go_terms") or []
    kegg = workflow.get("kegg_pathways") or []
    docking = workflow.get("docking_results") or []
    md_cfg = workflow.get("md_config") or {}

    prompt = f"""
# Task
Write a **publication-ready** IMRAD manuscript in Markdown for a network-pharmacology
study of **{plant}** against **{disease}**.

# Style rules
- Follow this exact section order and use `##` for section titles:
  {', '.join(SECTION_ORDER)}
- Abstract: 250-300 words, structured (Background, Methods, Results, Conclusion).
- Materials and Methods must name real tools: IMPPAT (compounds), admet-ai,
  ChEMBL + BindingDB (target prediction), Open Targets (disease targets),
  STRING (PPI), CytoHubba (hub ranking, mention MCC / Degree / Betweenness /
  Closeness / Radiality / EPC / MNC / DMNC / Stress / Bottleneck), g:Profiler
  (GO), Enrichr KEGG_2021_Human, AutoDock Vina 1.2.3 + Meeko + OpenBabel + RDKit,
  GROMACS with amber99sb-ildn + TIP3P.
- Results section MUST include sub-headings for every module and cite specific
  numbers from the DATA block below. Prefer sentences like "Compound X exhibited
  a docking affinity of −7.4 kcal/mol against Y".
- Discussion: 4-6 paragraphs interpreting biological significance.
- References: at least 15 real, well-known citations, formatted Vancouver style.
- Supplementary Tables: describe (do NOT embed) supplementary tables with
  captions (they will be attached as CSV/XLSX separately).

# Data
```json
{{
  "plant": "{plant}",
  "disease": "{disease}",
  "selected_compounds_top20": {_kbytes(compounds[:20], cap=2000)},
  "intersecting_genes_top50": {_kbytes(intersect[:50], cap=1200)},
  "hub_ranking_top10": {_kbytes(hubs[:10], cap=1200)},
  "top_go_terms": {_kbytes(go[:10], cap=1500)},
  "top_kegg_pathways": {_kbytes(kegg[:10], cap=1500)},
  "docking_summary_top10": {_kbytes(_slim_docking(docking), cap=2500)},
  "md_config": {_kbytes(md_cfg, cap=600)}
}}
```
Return **only Markdown**, no code fences, no meta commentary.
""".strip()

    # Safety net: if the assembled prompt is still huge, chop the data block.
    # Groq llama-3.3-70b context = ~32k tokens ≈ 120k chars; keep us well under.
    MAX_PROMPT_CHARS = 24000
    if len(prompt) > MAX_PROMPT_CHARS:
        prompt = prompt[:MAX_PROMPT_CHARS] + "\n… (data truncated to fit context window)\n\nReturn only Markdown."
    return prompt


async def generate_report(workflow: Dict[str, Any],
                          model: str = "claude-sonnet-4-5-20250929") -> Dict[str, Any]:
    prompt = _build_prompt(workflow)
    session_id = uuid.uuid4().hex
    logger.info(f"report/generate: prompt length = {len(prompt)} chars, session={session_id[:8]}")

    # ── Try Groq first (fastest + cheapest scientific writer) ──────────────
    try:
        import llm_groq
        if llm_groq.is_configured():
            # 4000 output tokens is enough for a full IMRAD manuscript. 45 s
            # upstream timeout keeps us well inside the Cloudflare ingress
            # timeout (~60 s) so we always surface a clean HTTP 500 with JSON
            # detail instead of a Cloudflare 502/520 error page.
            md = await llm_groq.chat_completion(
                [{"role": "system", "content": SYSTEM_PROMPT},
                 {"role": "user",   "content": prompt}],
                temperature=0.3,
                max_tokens=4000,
                timeout=45.0,
            )
            return {
                "markdown": md.strip(),
                "meta": {
                    "generated_at": datetime.now(timezone.utc).isoformat(),
                    "model": f"groq/{os.environ.get('GROQ_MODEL', 'llama-3.3-70b-versatile')}",
                    "plant": workflow.get("plant_name"),
                    "disease": workflow.get("disease_name"),
                    "session_id": session_id,
                },
            }
    except Exception as e:
        # Includes 413 Payload Too Large (workflow too big even after slimming),
        # 401 (bad key), 429 (rate-limited), timeouts, etc. Log + fall through.
        logger.warning(f"Groq scientific writer failed, falling back to Emergent: {e}")

    # ── Fallback: Emergent LLM key (Claude Sonnet) ──────────────
    api_key = os.environ.get("EMERGENT_LLM_KEY")
    if not api_key:
        return {"error": "No LLM configured (set GROQ_API_KEY or EMERGENT_LLM_KEY)",
                "markdown": "", "meta": {}}
    chat = LlmChat(api_key=api_key, session_id=session_id,
                   system_message=SYSTEM_PROMPT).with_model("anthropic", model)
    msg = UserMessage(text=prompt)
    try:
        # Emergent LLM's chat SDK doesn't expose a timeout kwarg — wrap with
        # asyncio.wait_for so a stuck upstream can't tie up the FastAPI worker
        # long enough for Cloudflare to return a 5xx. 45 s keeps us inside
        # the ingress timeout so we always return a clean HTTP 500 JSON.
        response = await asyncio.wait_for(chat.send_message(msg), timeout=45.0)
    except asyncio.TimeoutError:
        logger.error("Emergent LLM fallback timed out after 45 s")
        return {"error": "Report generation timed out. Try shrinking the workflow inputs "
                        "or splitting the run into fewer compounds/targets.",
                "markdown": "", "meta": {}}
    except Exception as e:
        logger.exception("LLM call failed")
        return {"error": f"Report generation failed: {e}", "markdown": "", "meta": {}}
    md = str(response).strip()
    return {
        "markdown": md,
        "meta": {
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "model": f"anthropic/{model}",
            "plant": workflow.get("plant_name"),
            "disease": workflow.get("disease_name"),
            "session_id": session_id,
        },
    }
    return {
        "markdown": md,
        "meta": {
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "model": f"anthropic/{model}",
            "plant": workflow.get("plant_name"),
            "disease": workflow.get("disease_name"),
            "session_id": session_id,
        },
    }


# ─────────────────────────── Export helpers ──────────────────────────────
def markdown_to_html(md: str, title: str = "PhytoNet AI — Research Report") -> str:
    try:
        import markdown as md_lib
        body = md_lib.markdown(md, extensions=["tables", "toc", "fenced_code"])
    except Exception:
        # Fallback: naive line-break rendering
        body = "<pre>" + md.replace("<", "&lt;") + "</pre>"
    return f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8"><title>{title}</title>
<style>
body {{ font-family: 'Georgia', 'Times New Roman', serif; max-width: 820px; margin: 40px auto; padding: 0 24px; color: #0B0B18; line-height: 1.65; }}
h1, h2 {{ font-family: 'Inter', 'Helvetica', sans-serif; color: #5139ED; }}
h2 {{ margin-top: 2.5rem; border-bottom: 1px solid #E7E7F3; padding-bottom: 4px; }}
h3 {{ margin-top: 1.5rem; color: #0B0B18; }}
table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; }}
th, td {{ border: 1px solid #E7E7F3; padding: 6px 10px; text-align: left; font-size: 13px; }}
th {{ background: #FAFAFF; }}
code {{ background: #F1F1FA; padding: 1px 4px; border-radius: 3px; }}
.footer {{ margin-top: 3rem; padding-top: 1rem; border-top: 1px solid #E7E7F3; color: #94A3B8; font-size: 12px; }}
</style></head>
<body>{body}
<p class="footer">Generated by PhytoNet AI · {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}</p>
</body></html>"""


def html_to_pdf(html: str) -> bytes:
    """Very light HTML-to-PDF via reportlab (text-only fallback if weasyprint absent)."""
    try:
        from weasyprint import HTML
        return HTML(string=html).write_pdf()
    except Exception:
        # Fallback: build a simple PDF via reportlab
        import io, re
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import LETTER
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=LETTER)
        w, h = LETTER
        text_only = re.sub(r"<[^>]+>", "", html)
        y = h - 60
        c.setFont("Times-Roman", 10)
        for line in text_only.splitlines():
            for chunk in [line[i:i + 100] for i in range(0, len(line) or 1, 100)]:
                if y < 60:
                    c.showPage(); c.setFont("Times-Roman", 10); y = h - 60
                c.drawString(60, y, chunk)
                y -= 14
        c.save()
        return buf.getvalue()


_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def _add_inline_runs(paragraph, text: str):
    """Parse a line of Markdown inline formatting (bold **, italic *, code `, links) into DOCX runs."""
    import re as _re
    parts = _INLINE_RE.split(text) if text else [text]
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2 and not part.startswith("**"):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"
        else:
            m = _re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                paragraph.add_run(m.group(1))  # DOCX hyperlinks require deeper OOXML — text-only
            else:
                paragraph.add_run(part)


def markdown_to_docx(md: str, title: Optional[str] = None) -> bytes:
    """Markdown → DOCX via python-docx.

    Supports: headings (#..####), bullet/numbered lists, blockquotes, tables
    (pipe syntax), horizontal rules, inline bold/italic/code.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)

    if title:
        h = doc.add_heading(title, level=0)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x51, 0x39, 0xED)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # Table detection: pipe-separated header + separator + body rows
        if (
            line.startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\s*\|[-:\s|]+\|\s*$", lines[i + 1])
        ):
            header = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            body = []
            while i < len(lines) and lines[i].startswith("|"):
                body.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            tbl = doc.add_table(rows=1 + len(body), cols=len(header))
            tbl.style = "Light Grid Accent 1"
            for c, cell in enumerate(header):
                p = tbl.rows[0].cells[c].paragraphs[0]
                r = p.add_run(cell); r.bold = True
            for ri, row in enumerate(body):
                for c, val in enumerate(row[: len(header)]):
                    p = tbl.rows[ri + 1].cells[c].paragraphs[0]
                    _add_inline_runs(p, val)
            continue

        if not line:
            doc.add_paragraph("")
        elif re.match(r"^-{3,}$", line) or re.match(r"^\*{3,}$", line):
            doc.add_paragraph("─" * 30)
        elif line.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            _add_inline_runs(p, line.lstrip("> ").strip())
        else:
            m = re.match(r"^(#{1,4})\s+(.*)$", line)
            if m:
                lvl = min(len(m.group(1)), 4)
                h = doc.add_heading(m.group(2), level=lvl)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x51, 0x39, 0xED) if lvl <= 2 else RGBColor(0x0B, 0x0B, 0x18)
            elif line.lstrip().startswith(("- ", "* ")):
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, line.lstrip()[2:])
            elif re.match(r"^\d+\.\s", line.lstrip()):
                p = doc.add_paragraph(style="List Number")
                _add_inline_runs(p, re.sub(r"^\d+\.\s", "", line.lstrip()))
            else:
                p = doc.add_paragraph()
                _add_inline_runs(p, line)
        i += 1

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()
